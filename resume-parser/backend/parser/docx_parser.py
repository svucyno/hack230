"""
DOCX text extraction using python-docx.
Extracts paragraph text and table cell content.
"""

import io
from pathlib import Path
from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file.

    Args:
        file_path: Absolute or relative path to the .docx file.

    Returns:
        A single string containing all extracted text.

    Raises:
        FileNotFoundError: If the file doesn't exist.
        RuntimeError: If python-docx fails to open the file.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX '{file_path}': {exc}") from exc

    return _extract_from_document(doc)


def extract_text_from_docx_bytes(file_bytes: bytes, filename: str = "upload.docx") -> str:
    """
    Extract text directly from DOCX bytes (for FastAPI UploadFile usage).

    Args:
        file_bytes: Raw bytes of the DOCX file.
        filename:   Name used only for error messages.

    Returns:
        Extracted text string.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise RuntimeError(f"Failed to parse DOCX bytes for '{filename}': {exc}") from exc

    return _extract_from_document(doc)


def _extract_from_document(doc: Document) -> str:
    """
    Internal helper: pull text from paragraphs and tables.
    """
    lines = []

    # Paragraphs (main body text)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)

    # Tables (some resumes put info in tables)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                lines.append(" | ".join(row_texts))

    return "\n".join(lines).strip()
